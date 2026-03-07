#!/usr/bin/env python3
"""
DOCX Extraction Framework

Extracts text, tables, and embedded images from DOCX files with
positional metadata for downstream parsing.
"""

import json
import os
import sys
from docx import Document
from docx.oxml.ns import qn


def extract_images(doc: Document, output_dir: str, doc_name: str) -> list:
    """Extract embedded images from a DOCX, saving them to output_dir.
    Returns a list of dicts with image metadata and positional info."""
    images = []
    rels = doc.part.rels
    image_idx = 0

    for rel_id, rel in rels.items():
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            ext = os.path.splitext(rel.target_ref)[1] or '.png'
            fname = f"{doc_name}_img{image_idx}{ext}"
            fpath = os.path.join(output_dir, fname)
            with open(fpath, 'wb') as f:
                f.write(image_data)
            images.append({
                'index': image_idx,
                'rel_id': rel_id,
                'filename': fname,
                'path': fpath,
                'size_bytes': len(image_data),
            })
            image_idx += 1

    return images


def extract_tables(doc: Document) -> list:
    """Extract all tables from a DOCX as lists of rows (list of cells)."""
    tables = []
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip())
            rows.append(cells)
        tables.append({
            'index': table_idx,
            'rows': rows,
            'num_rows': len(rows),
            'num_cols': len(rows[0]) if rows else 0,
        })
    return tables


def extract_paragraphs(doc: Document) -> list:
    """Extract all paragraphs with their style names."""
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else 'Normal'

        # Check for bold/italic runs
        has_bold = any(run.bold for run in para.runs if run.bold)
        has_italic = any(run.italic for run in para.runs if run.italic)

        paragraphs.append({
            'text': text,
            'style': style_name,
            'has_bold': has_bold,
            'has_italic': has_italic,
        })
    return paragraphs


def extract_body_elements(doc: Document) -> list:
    """Extract body elements in document order (paragraphs and tables interleaved).
    This preserves the original ordering which is critical for associating
    content with the correct year level/section."""
    elements = []
    table_idx = 0
    para_idx = 0

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            if para_idx < len(doc.paragraphs):
                para = doc.paragraphs[para_idx]
                text = para.text.strip()
                style_name = para.style.name if para.style else 'Normal'
                if text:
                    elements.append({
                        'type': 'paragraph',
                        'text': text,
                        'style': style_name,
                        'index': para_idx,
                    })
                para_idx += 1
        elif tag == 'tbl':
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                elements.append({
                    'type': 'table',
                    'rows': rows,
                    'num_rows': len(rows),
                    'num_cols': len(rows[0]) if rows else 0,
                    'index': table_idx,
                })
                table_idx += 1

    return elements


def extract_docx(filepath: str, images_dir: str = None) -> dict:
    """Full extraction of a DOCX file.

    Returns a dict with:
      - filename: basename of the file
      - paragraphs: list of paragraph dicts
      - tables: list of table dicts
      - body_elements: ordered list of paragraphs and tables as they appear
      - images: list of extracted image metadata (if images_dir provided)
    """
    doc = Document(filepath)
    basename = os.path.splitext(os.path.basename(filepath))[0]

    result = {
        'filename': os.path.basename(filepath),
        'paragraphs': extract_paragraphs(doc),
        'tables': extract_tables(doc),
        'body_elements': extract_body_elements(doc),
        'images': [],
    }

    if images_dir:
        os.makedirs(images_dir, exist_ok=True)
        result['images'] = extract_images(doc, images_dir, basename)

    return result


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: extract_docx.py <file.docx> [images_output_dir]")
        sys.exit(1)

    filepath = sys.argv[1]
    images_dir = sys.argv[2] if len(sys.argv) > 2 else None
    result = extract_docx(filepath, images_dir)
    print(json.dumps(result, indent=2, ensure_ascii=False))
