#!/usr/bin/env python3
"""
Parser for SCSA judging standards DOCX files.

Extracts A-E grade descriptors from SCSA judging standards documents.
These documents use nested tables within outer layout tables.  Each nested
grade table has headers (A–E) in row 0, optional section-header rows where
every cell repeats the same text, and content rows with strand names in
column 0 and grade descriptors in columns 1–5.
"""

import os
import re
import sys
import json
import argparse

from docx import Document


GRADES = ["A", "B", "C", "D", "E"]

# Normalised filename area → database learning_area_id.
AREA_MAP = {
    "science": "science",
    "english": "english",
    "mathematics": "mathematics",
    "technologies_digital_technologies": "tech-digital",
    "technologies_design_and_technologies": "tech-design",
    "health_and_physical_education_health_education": "hpe",
    "health_and_physical_education_physical_education": "hpe",
    "hass_civics_citizenship": "hass-civics",
    "hass_economics_business": "hass-economics",
    "hass_geography": "hass-geography",
    "hass_history": "hass-history",
}


def _normalise(s: str) -> str:
    """Lower-case and collapse all separators to underscores."""
    return s.lower().replace("-", "_").replace(" ", "_")


def _parse_filename(filename: str) -> tuple[str, str, bool]:
    """Return (year_level_id, learning_area_id, is_best_effort) from *filename*."""
    stem = re.sub(r"\.docx$", "", filename, flags=re.IGNORECASE)

    is_best_effort = bool(re.search(r"_20\d{2}$", stem))
    stem = re.sub(r"_20\d{2}$", "", stem)

    year_match = re.search(r"Year[_-](PP|\d+)", stem, re.IGNORECASE)
    if not year_match:
        raise ValueError(f"Cannot extract year level from: {filename}")
    year_raw = year_match.group(1)
    year_level_id = "PP" if year_raw.upper() == "PP" else f"Y{year_raw}"

    if stem.upper().startswith("HASS"):
        area_match = re.match(r"(HASS[_\-][\w-]+?)[_-]Year", stem, re.IGNORECASE)
        area_raw = area_match.group(1) if area_match else "unknown"
    else:
        area_match = re.search(
            r"Year[_-](?:PP|\d+)[_-](.+?)[_-]Judging", stem, re.IGNORECASE
        )
        area_raw = area_match.group(1) if area_match else "unknown"

    area_norm = _normalise(area_raw)
    learning_area_id = AREA_MAP.get(area_norm, "unknown")
    if learning_area_id == "unknown":
        raise ValueError(
            f"Unknown learning area '{area_raw}' (normalised: '{area_norm}') "
            f"from: {filename}"
        )

    return year_level_id, learning_area_id, is_best_effort


# ------------------------------------------------------------------
# Nested-table helpers
# ------------------------------------------------------------------

def _is_grade_header_row(row) -> bool:
    """True when the row contains the A–E grade column headers."""
    if len(row.cells) < 6:
        return False
    return any("xcellent" in c.text for c in row.cells)


def _is_section_header_row(row) -> bool:
    """True for rows where every cell repeats the same label (e.g. 'Science Understanding')."""
    first = row.cells[0].text.strip()
    if not first:
        return False
    return all(c.text.strip() == first for c in row.cells)


def _parse_grade_table(table) -> list[dict]:
    """Parse a single nested grade table into strand / grade / descriptor dicts."""
    if len(table.rows) < 2 or len(table.columns) < 6:
        return []
    if not _is_grade_header_row(table.rows[0]):
        return []

    strands: list[tuple[str, dict[str, list[str]]]] = []
    current_strand: str | None = None
    current_descs: dict[str, list[str]] | None = None

    for i in range(1, len(table.rows)):
        row = table.rows[i]

        if _is_section_header_row(row):
            continue

        col0 = row.cells[0].text.strip()

        if col0:
            # New strand — flush the previous one
            if current_strand is not None:
                strands.append((current_strand, current_descs))
            current_strand = col0
            current_descs = {g: [] for g in GRADES}
            for gi, grade in enumerate(GRADES):
                text = row.cells[gi + 1].text.strip()
                if text:
                    current_descs[grade].append(text)
        elif current_strand is not None:
            # Continuation row — append extra descriptor text
            for gi, grade in enumerate(GRADES):
                text = row.cells[gi + 1].text.strip()
                if text:
                    current_descs[grade].append(text)

    if current_strand is not None:
        strands.append((current_strand, current_descs))

    records: list[dict] = []
    for strand_name, descs in strands:
        for grade in GRADES:
            if descs[grade]:
                records.append(
                    {
                        "strand": strand_name,
                        "grade": grade,
                        "descriptor_text": "\n".join(descs[grade]),
                    }
                )
    return records


def _find_grade_tables(doc: Document) -> list:
    """Collect every unique nested table that carries A–E grade headers."""
    seen: set[int] = set()
    grade_tables = []

    for table in doc.tables:
        for row in table.rows:
            # Only inspect cell 0 — other cells duplicate content for
            # landscape printing.
            cell = row.cells[0]
            for nested in cell.tables:
                tid = id(nested)
                if tid in seen:
                    continue
                seen.add(tid)
                if nested.rows and _is_grade_header_row(nested.rows[0]):
                    grade_tables.append(nested)

    return grade_tables


# ------------------------------------------------------------------
# Public API
# ------------------------------------------------------------------

def parse_judging_standards(filepath: str) -> list[dict]:
    """
    Parse a SCSA judging standards DOCX file into grade descriptor records.

    Args:
        filepath: Path to the judging standards DOCX file.

    Returns:
        List of dicts with keys: learning_area_id, year_level_id, strand,
        grade, descriptor_text, source_document, is_best_effort.
    """
    filename = os.path.basename(filepath)
    year_level_id, learning_area_id, is_best_effort = _parse_filename(filename)

    doc = Document(filepath)
    grade_tables = _find_grade_tables(doc)

    results: list[dict] = []
    for gt in grade_tables:
        for record in _parse_grade_table(gt):
            record["learning_area_id"] = learning_area_id
            record["year_level_id"] = year_level_id
            record["source_document"] = filename
            record["is_best_effort"] = is_best_effort
            results.append(record)

    return results


def main():
    parser = argparse.ArgumentParser(
        description="Parse SCSA judging standards DOCX files"
    )
    parser.add_argument("filepath", help="Path to judging standards DOCX file")
    parser.add_argument("--json", action="store_true", help="Output as JSON")
    args = parser.parse_args()

    results = parse_judging_standards(args.filepath)

    if args.json:
        print(json.dumps(results, indent=2))
    else:
        print(
            f"Parsed {len(results)} grade descriptors from "
            f"{os.path.basename(args.filepath)}"
        )
        for r in results:
            print(f"  [{r['grade']}] {r['strand']}: {r['descriptor_text'][:80]}...")


if __name__ == "__main__":
    main()
